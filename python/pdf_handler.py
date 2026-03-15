"""PDF Handler - Handles all PDF operations."""
import sys
import json
import os
import io
import time
import re
import itertools
import argparse
import tempfile
import base64
from pathlib import Path
from path_utils import (
    resolve_input_path,
    resolve_output_dir,
    resolve_output_file,
    default_single_output,
    default_first_input_output,
    create_unique_child_dir,
)


def merge_pdfs(input_files: list, output_file: str = "") -> dict:
    """Merge multiple PDFs into one."""
    from pypdf import PdfWriter
    writer = PdfWriter()
    try:
        if not input_files:
            return {"success": False, "error": "No input files provided"}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_first_input_output(input_files, "_merged", ".pdf", __file__)
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        from pypdf import PdfReader
        for f in input_files:
            resolved = resolve_input_path(f, __file__)
            if not os.path.exists(resolved):
                return {"success": False, "error": f"Input file not found: {f}. Please select the file using native file picker so absolute path is available."}
            reader = PdfReader(resolved)
            for page in reader.pages:
                writer.add_page(page)
        with open(output_file, "wb") as out:
            writer.write(out)
        return {"success": True, "output": output_file}
    except Exception as e:
        return {"success": False, "error": str(e)}


def split_pdf(input_file: str, output_dir: str = "", ranges: list = None) -> dict:
    """Split PDF into individual pages or by ranges."""
    from pypdf import PdfReader, PdfWriter
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        if output_dir and str(output_dir).strip():
            output_dir = resolve_output_dir(output_dir, __file__)
        else:
            output_dir = create_unique_child_dir(input_file, "split")
        reader = PdfReader(input_file)
        total_pages = len(reader.pages)
        os.makedirs(output_dir, exist_ok=True)
        stem = Path(input_file).stem
        outputs = []

        if ranges:
            for i, (start, end) in enumerate(ranges):
                writer = PdfWriter()
                for p in range(start - 1, min(end, total_pages)):
                    writer.add_page(reader.pages[p])
                out_path = os.path.join(output_dir, f"{stem}_part{i+1}.pdf")
                with open(out_path, "wb") as f:
                    writer.write(f)
                outputs.append(out_path)
        else:
            for i, page in enumerate(reader.pages):
                writer = PdfWriter()
                writer.add_page(page)
                out_path = os.path.join(output_dir, f"{stem}_page{i+1}.pdf")
                with open(out_path, "wb") as f:
                    writer.write(f)
                outputs.append(out_path)

        return {"success": True, "outputs": outputs, "total_pages": total_pages}
    except Exception as e:
        return {"success": False, "error": str(e)}


def encrypt_pdf(input_file: str, output_file: str = "", password: str = "") -> dict:
    """Encrypt PDF with password."""
    from pypdf import PdfReader, PdfWriter
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "_encrypted", ".pdf")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        reader = PdfReader(input_file)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(password)
        with open(output_file, "wb") as f:
            writer.write(f)
        return {"success": True, "output": output_file}
    except Exception as e:
        return {"success": False, "error": str(e)}


def decrypt_pdf(input_file: str, output_file: str = "", password: str = "") -> dict:
    """Decrypt password-protected PDF."""
    from pypdf import PdfReader, PdfWriter
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "_decrypted", ".pdf")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        reader = PdfReader(input_file)
        if reader.is_encrypted:
            reader.decrypt(password)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(output_file, "wb") as f:
            writer.write(f)
        return {"success": True, "output": output_file}
    except Exception as e:
        return {"success": False, "error": str(e)}


def compress_pdf(input_file: str, output_file: str = "") -> dict:
    """Compress PDF by removing redundant data."""
    from pypdf import PdfReader, PdfWriter
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "_compressed", ".pdf")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        reader = PdfReader(input_file)
        writer = PdfWriter()
        # pypdf 5.x can require pages to belong to a writer before compression.
        # Add first, then compress the writer-owned page object.
        for page in reader.pages:
            writer.add_page(page)
        for page in writer.pages:
            page.compress_content_streams()
        with open(output_file, "wb") as f:
            writer.write(f)
        original_size = os.path.getsize(input_file)
        compressed_size = os.path.getsize(output_file)
        return {
            "success": True,
            "output": output_file,
            "original_size": original_size,
            "compressed_size": compressed_size,
            "ratio": round((1 - compressed_size / original_size) * 100, 1)
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def add_watermark(input_file: str, output_file: str = "", text: str = "",
                  font_size: int = 40, opacity: float = 0.3,
                  color: str = "gray", image_path: str = "",
                  image_scale: float = 0.4) -> dict:
    """Add text or image watermark to each PDF page."""
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    import io
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "_watermarked", ".pdf")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        image_path = resolve_input_path(image_path, __file__) if image_path and str(image_path).strip() else ""
        if not text and not image_path:
            return {"success": False, "error": "Please provide watermark text or image."}
        if image_path and not os.path.exists(image_path):
            return {"success": False, "error": f"Watermark image not found: {image_path}"}

        reader = PdfReader(input_file)
        writer = PdfWriter()
        color_map = {
            "gray": colors.gray,
            "red": colors.red,
            "blue": colors.blue,
            "black": colors.black,
        }
        fill_color = color_map.get(color, colors.gray)

        for page in reader.pages:
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=(w, h))
            c.saveState()
            c.setFillColor(fill_color, alpha=opacity)
            c.translate(w / 2, h / 2)
            c.rotate(45)

            if image_path:
                img = ImageReader(image_path)
                img_w, img_h = img.getSize()
                safe_scale = max(0.1, min(float(image_scale), 1.2))
                target_w = w * safe_scale
                target_h = target_w * (img_h / img_w)
                if hasattr(c, "setFillAlpha"):
                    c.setFillAlpha(opacity)
                c.drawImage(
                    img,
                    -target_w / 2,
                    -target_h / 2,
                    width=target_w,
                    height=target_h,
                    mask='auto',
                    preserveAspectRatio=True,
                )

            if text:
                c.setFont("Helvetica", font_size)
                c.drawCentredString(0, 0, text)
            c.restoreState()
            c.save()
            packet.seek(0)
            from pypdf import PdfReader as PR
            watermark_page = PR(packet).pages[0]
            page.merge_page(watermark_page)
            writer.add_page(page)

        with open(output_file, "wb") as f:
            writer.write(f)
        return {"success": True, "output": output_file}
    except Exception as e:
        return {"success": False, "error": str(e)}


def add_page_numbers(input_file: str, output_file: str = "",
                     position: str = "bottom-center",
                     start: int = 1, font_size: int = 12) -> dict:
    """Add page numbers to PDF."""
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    import io
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "_numbered", ".pdf")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)
        reader = PdfReader(input_file)
        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            w = float(page.mediabox.width)
            h = float(page.mediabox.height)
            packet = io.BytesIO()
            c = canvas.Canvas(packet, pagesize=(w, h))
            num_text = str(i + start)
            c.setFont("Helvetica", font_size)

            if position == "bottom-center":
                c.drawCentredString(w / 2, 20, num_text)
            elif position == "bottom-right":
                c.drawRightString(w - 20, 20, num_text)
            elif position == "bottom-left":
                c.drawString(20, 20, num_text)
            elif position == "top-center":
                c.drawCentredString(w / 2, h - 20, num_text)

            c.save()
            packet.seek(0)
            from pypdf import PdfReader as PR
            num_page = PR(packet).pages[0]
            page.merge_page(num_page)
            writer.add_page(page)

        with open(output_file, "wb") as f:
            writer.write(f)
        return {"success": True, "output": output_file}
    except Exception as e:
        return {"success": False, "error": str(e)}


def preview_page_numbers(input_file: str,
                         position: str = "bottom-center",
                         start: int = 1,
                         font_size: int = 12) -> dict:
    """Generate a one-page PNG preview for page number settings."""
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}"}

        from PIL import Image, ImageDraw, ImageFont

        preview = None
        try:
            from pdf2image import convert_from_path
            pages = convert_from_path(input_file, first_page=1, last_page=1, dpi=120)
            if pages:
                preview = pages[0].convert("RGBA")
        except Exception:
            preview = None

        if preview is None:
            preview = Image.new("RGBA", (900, 1273), (255, 255, 255, 255))

        draw = ImageDraw.Draw(preview)
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", max(10, int(font_size * 2)))
        except Exception:
            font = ImageFont.load_default()

        text = str(start)
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]

        if position == "bottom-center":
            x, y = (preview.width - tw) // 2, preview.height - th - 24
        elif position == "bottom-right":
            x, y = preview.width - tw - 24, preview.height - th - 24
        elif position == "bottom-left":
            x, y = 24, preview.height - th - 24
        else:
            x, y = (preview.width - tw) // 2, 24

        draw.text((x, y), text, fill=(35, 35, 35, 255), font=font)

        out_path = os.path.join(tempfile.gettempdir(), f"dochub_preview_page_numbers_{int(time.time() * 1000)}.png")
        preview.convert("RGB").save(out_path, format="PNG")
        with open(out_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")

        return {"success": True, "output": out_path, "preview_data_url": f"data:image/png;base64,{b64}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def preview_watermark(input_file: str,
                      text: str = "",
                      font_size: int = 40,
                      opacity: float = 0.3,
                      color: str = "gray",
                      image_path: str = "",
                      image_scale: float = 0.4) -> dict:
    """Generate a one-page PNG preview for watermark settings."""
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}"}

        image_path = resolve_input_path(image_path, __file__) if image_path and str(image_path).strip() else ""
        if not text and not image_path:
            return {"success": False, "error": "Please provide watermark text or image."}
        if image_path and not os.path.exists(image_path):
            return {"success": False, "error": f"Watermark image not found: {image_path}"}

        from PIL import Image, ImageDraw, ImageFont

        preview = None
        try:
            from pdf2image import convert_from_path
            pages = convert_from_path(input_file, first_page=1, last_page=1, dpi=120)
            if pages:
                preview = pages[0].convert("RGBA")
        except Exception:
            preview = None

        if preview is None:
            preview = Image.new("RGBA", (900, 1273), (255, 255, 255, 255))

        overlay = Image.new("RGBA", preview.size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(overlay)
        alpha = int(max(0.05, min(float(opacity), 1.0)) * 255)

        color_map = {
            "gray": (107, 114, 128, alpha),
            "red": (220, 38, 38, alpha),
            "blue": (37, 99, 235, alpha),
            "black": (17, 24, 39, alpha),
        }
        fill = color_map.get(color, color_map["gray"])

        if image_path:
            wm = Image.open(image_path).convert("RGBA")
            safe_scale = max(0.1, min(float(image_scale), 1.2))
            target_w = max(60, int(preview.width * safe_scale))
            target_h = max(60, int(target_w * wm.height / max(1, wm.width)))
            wm = wm.resize((target_w, target_h))
            if alpha < 255:
                wm_alpha = wm.split()[3].point(lambda p: int(p * alpha / 255))
                wm.putalpha(wm_alpha)

            rot = wm.rotate(25, expand=True)
            x = (preview.width - rot.width) // 2
            y = (preview.height - rot.height) // 2
            overlay.alpha_composite(rot, (x, y))

        if text:
            try:
                font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", max(12, int(font_size * 2)))
            except Exception:
                font = ImageFont.load_default()
            bbox = draw.textbbox((0, 0), text, font=font)
            tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            tx = (preview.width - tw) // 2
            ty = (preview.height - th) // 2
            draw.text((tx, ty), text, fill=fill, font=font)
            overlay = overlay.rotate(25, center=(preview.width // 2, preview.height // 2), expand=False)

        result = Image.alpha_composite(preview, overlay)
        out_path = os.path.join(tempfile.gettempdir(), f"dochub_preview_watermark_{int(time.time() * 1000)}.png")
        result.convert("RGB").save(out_path, format="PNG")
        with open(out_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")

        return {"success": True, "output": out_path, "preview_data_url": f"data:image/png;base64,{b64}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def pdf_to_docx(input_file: str,
                output_file: str = "",
                ignore_page_error: bool = True,
                multi_processing: bool = False,
                cpu_count: int = 0,
                parse_lattice_table: bool = True,
                parse_stream_table: bool = True,
                delete_end_line_hyphen: bool = False,
                raw_exceptions: bool = False,
                provider: str = "local",
                credentials: dict = None,
                lang: str = "chi_sim+eng") -> dict:
    """Convert PDF to Word document using tunable pdf2docx settings."""
    from pdf2docx import Converter
    from pypdf import PdfReader

    def _to_bool(value, default):
        if isinstance(value, bool):
            return value
        if isinstance(value, str):
            text = value.strip().lower()
            if text in {"1", "true", "yes", "on"}:
                return True
            if text in {"0", "false", "no", "off"}:
                return False
        if isinstance(value, (int, float)):
            return bool(value)
        return default

    def _markdown_to_docx(md_text: str, output_path: str):
        from docx import Document

        doc = Document()
        for raw_line in (md_text or "").splitlines():
            line = raw_line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                title = line[level:].strip() or line
                doc.add_heading(title, level=min(max(level, 1), 4))
                continue

            if line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
                continue

            if re.match(r"^\d+\.\s+", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
                continue

            if line.startswith("!["):
                # Keep OCR markdown images out of plain docx text fallback.
                continue

            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            doc.add_paragraph(clean)

        doc.save(output_path)

    cv = None
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        if Path(input_file).suffix.lower() != ".pdf":
            return {"success": False, "error": "Input file must be a PDF"}

        # Quick preflight for clearer errors before heavy layout parsing.
        reader = PdfReader(input_file)
        if reader.is_encrypted:
            return {"success": False, "error": "Encrypted PDF is not supported for PDF to Word. Please decrypt first."}
        total_pages = len(reader.pages)
        if total_pages <= 0:
            return {"success": False, "error": "Empty PDF"}

        output_file = resolve_output_file(output_file, __file__) if output_file and str(output_file).strip() else default_single_output(input_file, "", ".docx")
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)

        safe_multi_processing = _to_bool(multi_processing, False)
        safe_ignore_page_error = _to_bool(ignore_page_error, True)
        safe_parse_lattice_table = _to_bool(parse_lattice_table, True)
        safe_parse_stream_table = _to_bool(parse_stream_table, True)
        safe_delete_end_line_hyphen = _to_bool(delete_end_line_hyphen, False)
        safe_raw_exceptions = _to_bool(raw_exceptions, False)

        cpu_max = max(1, os.cpu_count() or 1)
        try:
            safe_cpu_count = int(cpu_count)
        except Exception:
            safe_cpu_count = 0

        warnings = []
        credentials = credentials or {}
        provider_norm = str(provider or "local").strip().lower()
        if provider_norm not in {"auto", "local", "glm"}:
            warnings.append(f"unknown provider '{provider_norm}', fallback to local")
            provider_norm = "local"
        api_key = credentials.get("api_key") if isinstance(credentials, dict) else None

        use_glm = provider_norm in {"auto", "glm"} and bool(api_key)
        if provider_norm == "glm" and not api_key:
            return {"success": False, "error": "GLM provider selected but credentials.api_key is missing"}

        if safe_cpu_count < 0:
            warnings.append("cpu_count is negative and has been reset to 0")
            safe_cpu_count = 0
        if safe_cpu_count > cpu_max:
            warnings.append(f"cpu_count exceeds available CPUs ({cpu_max}) and has been clamped")
            safe_cpu_count = cpu_max
        if not safe_multi_processing and safe_cpu_count > 0:
            warnings.append("cpu_count is ignored when multi_processing is disabled")

        if use_glm:
            from ocr_handler import ocr_pdf

            glm_res = ocr_pdf(
                pdf_path=input_file,
                lang=lang,
                provider="glm",
                credentials={"api_key": api_key},
            )
            if glm_res.get("success") and (glm_res.get("text") or "").strip():
                _markdown_to_docx(glm_res.get("text", ""), output_file)
                result = {
                    "success": True,
                    "output": output_file,
                    "source": "glm",
                    "pages": glm_res.get("pages"),
                    "config_applied": {
                        "provider": provider_norm,
                        "lang": lang,
                    },
                }
                if warnings:
                    result["warnings"] = warnings
                return result

            if provider_norm == "glm":
                return {"success": False, "error": glm_res.get("error", "GLM OCR failed"), "source": "glm"}

            warnings.append("GLM OCR failed in auto mode, fallback to local pdf2docx")

        cv = Converter(input_file)
        convert_settings = cv.default_settings.copy()
        convert_settings.update({
            "ignore_page_error": safe_ignore_page_error,
            "multi_processing": safe_multi_processing,
            "cpu_count": safe_cpu_count,
            "parse_lattice_table": safe_parse_lattice_table,
            "parse_stream_table": safe_parse_stream_table,
            "delete_end_line_hyphen": safe_delete_end_line_hyphen,
            "raw_exceptions": safe_raw_exceptions,
        })

        cv.convert(output_file, start=0, end=None, **convert_settings)

        result = {
            "success": True,
            "output": output_file,
            "source": "pdf2docx",
            "config_applied": {
                "provider": provider_norm,
                "ignore_page_error": convert_settings["ignore_page_error"],
                "multi_processing": convert_settings["multi_processing"],
                "cpu_count": convert_settings["cpu_count"],
                "parse_lattice_table": convert_settings["parse_lattice_table"],
                "parse_stream_table": convert_settings["parse_stream_table"],
                "delete_end_line_hyphen": convert_settings["delete_end_line_hyphen"],
                "raw_exceptions": convert_settings["raw_exceptions"],
            },
        }
        if warnings:
            result["warnings"] = warnings
        return result
    except Exception as e:
        return {"success": False, "error": str(e)}
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:
                pass


def preview_pages(input_file: str,
                  dpi: int = 90,
                  max_width: int = 260) -> dict:
    """Generate thumbnail images for all pages in a PDF."""
    from pypdf import PdfReader
    from PIL import Image, ImageDraw, ImageFont
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}"}

        reader = PdfReader(input_file)
        total_pages = len(reader.pages)
        if total_pages <= 0:
            return {"success": False, "error": "Empty PDF"}

        safe_dpi = max(60, min(int(dpi), 180))
        safe_max_width = max(140, min(int(max_width), 480))

        out_dir = os.path.join(
            tempfile.gettempdir(),
            f"dochub_preview_pages_{int(time.time() * 1000)}"
        )
        os.makedirs(out_dir, exist_ok=True)

        previews = []

        def _append_preview(page_no: int, image_path: str):
            with open(image_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")
            previews.append({
                "page": page_no,
                "image": image_path,
                "preview_data_url": f"data:image/png;base64,{b64}",
            })

        try:
            from pdf2image import convert_from_path
            images = convert_from_path(input_file, dpi=safe_dpi, fmt="png")
            for idx, img in enumerate(images, start=1):
                if img.width > safe_max_width:
                    target_h = max(1, int(img.height * safe_max_width / img.width))
                    img = img.resize((safe_max_width, target_h), Image.Resampling.LANCZOS)

                out_path = os.path.join(out_dir, f"page_{idx}.png")
                img.save(out_path, format="PNG", optimize=True)
                _append_preview(idx, out_path)
        except Exception:
            # Fallback to simple placeholders when PDF rasterization is unavailable.
            for idx in range(1, total_pages + 1):
                img = Image.new("RGB", (safe_max_width, int(safe_max_width * 1.4)), (255, 255, 255))
                draw = ImageDraw.Draw(img)
                draw.rectangle((0, 0, img.width - 1, img.height - 1), outline=(218, 223, 234), width=2)

                try:
                    font = ImageFont.truetype(
                        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                        max(14, safe_max_width // 8),
                    )
                except Exception:
                    font = ImageFont.load_default()

                text = f"Page {idx}"
                bbox = draw.textbbox((0, 0), text, font=font)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                draw.text(
                    ((img.width - tw) // 2, (img.height - th) // 2),
                    text,
                    fill=(44, 62, 90),
                    font=font,
                )

                out_path = os.path.join(out_dir, f"page_{idx}.png")
                img.save(out_path, format="PNG", optimize=True)
                _append_preview(idx, out_path)

        return {
            "success": True,
            "total_pages": total_pages,
            "previews": previews,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def reorder_pdf_pages(input_file: str,
                      page_order: list,
                      output_file: str = "") -> dict:
    """Reorder PDF pages according to a 1-based page index list."""
    from pypdf import PdfReader, PdfWriter
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}"}

        reader = PdfReader(input_file)
        total_pages = len(reader.pages)
        if total_pages <= 0:
            return {"success": False, "error": "Empty PDF"}

        if not isinstance(page_order, list) or not page_order:
            return {"success": False, "error": "page_order must be a non-empty list"}

        try:
            normalized = [int(p) for p in page_order]
        except Exception:
            return {"success": False, "error": "page_order must contain integers"}

        expected = list(range(1, total_pages + 1))
        if sorted(normalized) != expected:
            return {
                "success": False,
                "error": f"page_order must contain each page exactly once (1..{total_pages})",
            }

        output_file = (
            resolve_output_file(output_file, __file__)
            if output_file and str(output_file).strip()
            else default_single_output(input_file, "_reordered", ".pdf")
        )
        Path(output_file).parent.mkdir(parents=True, exist_ok=True)

        writer = PdfWriter()
        for p in normalized:
            writer.add_page(reader.pages[p - 1])

        with open(output_file, "wb") as f:
            writer.write(f)

        return {
            "success": True,
            "output": output_file,
            "total_pages": total_pages,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def get_pdf_info(input_file: str) -> dict:
    """Get PDF metadata and info."""
    from pypdf import PdfReader
    try:
        input_file = resolve_input_path(input_file, __file__)
        if not os.path.exists(input_file):
            return {"success": False, "error": f"Input file not found: {input_file}. Please select the file using native file picker so absolute path is available."}
        reader = PdfReader(input_file)
        meta = reader.metadata or {}
        return {
            "success": True,
            "pages": len(reader.pages),
            "encrypted": reader.is_encrypted,
            "title": meta.get("/Title", ""),
            "author": meta.get("/Author", ""),
            "size": os.path.getsize(input_file),
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ── Brute Force ───────────────────────────────────────────────────────────────

CHARSETS = {
    "digits":  "0123456789",
    "lower":   "abcdefghijklmnopqrstuvwxyz",
    "upper":   "ABCDEFGHIJKLMNOPQRSTUVWXYZ",
    "symbols": "!@#$%^&*()-_=+[]{}|;:,.<>?/",
}

# Module-level reader for each worker process (set by pool initializer)
_bf_reader = None


def _bf_init(shm_name: str, shm_size: int):
    """Pool worker initializer: load PDF once per process via shared memory."""
    global _bf_reader
    from multiprocessing.shared_memory import SharedMemory
    from pypdf import PdfReader
    shm = SharedMemory(name=shm_name)
    data = bytes(shm.buf[:shm_size])
    shm.close()
    _bf_reader = PdfReader(io.BytesIO(data))


def _bf_try_chunk(passwords: list):
    """Try a list of passwords; return the correct one or None."""
    for pwd in passwords:
        try:
            if _bf_reader.decrypt(pwd) != 0:
                return pwd
        except Exception:
            pass
    return None


def _chunk_gen(iterable, size: int):
    """Yield successive chunks of given size from an iterable."""
    it = iter(iterable)
    while True:
        chunk = list(itertools.islice(it, size))
        if not chunk:
            break
        yield chunk


def bruteforce_pdf(pdf_path: str,
                   mode: str = "charset",
                   charset_keys: list = None,
                   custom_charset: str = "",
                   min_len: int = 1,
                   max_len: int = 4,
                   dict_path: str = None,
                   num_workers: int = 0) -> dict:
    """
    Brute-force PDF password using multiprocessing (bypasses GIL for true
    parallelism). Hashlib/OpenSSL underneath uses CPU SIMD intrinsics for MD5.

    Progress is streamed to stdout as JSON lines so Rust can emit Tauri events:
      {"type":"progress","tried":N,"total":T,"speed":S,"eta":E,"elapsed":X}
      {"type":"found","password":"...","tried":N,"elapsed":X}
      {"type":"done","found":false,"tried":N,"elapsed":X}
      {"type":"error","error":"..."}
    """
    import multiprocessing
    import itertools
    from multiprocessing.shared_memory import SharedMemory
    from pypdf import PdfReader

    # ── Validate ──────────────────────────────────────────────────────────────
    pdf_path = resolve_input_path(pdf_path, __file__)
    if not os.path.exists(pdf_path):
        _bf_print("error", error="File not found")
        return {"success": False, "error": "File not found"}

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    probe = PdfReader(io.BytesIO(pdf_bytes))
    if not probe.is_encrypted:
        _bf_print("error", error="PDF is not encrypted")
        return {"success": False, "error": "PDF is not encrypted"}

    n_workers = num_workers if num_workers > 0 else multiprocessing.cpu_count()
    chunk_size = max(200, 1000 // n_workers)

    # ── Build charset / wordlist ──────────────────────────────────────────────
    if mode == "charset":
        keys = charset_keys or ["digits"]
        raw = "".join(CHARSETS.get(k, k) for k in keys) + (custom_charset or "")
        # deduplicate preserving order
        seen_chars: set = set()
        charset = "".join(c for c in raw if not (c in seen_chars or seen_chars.add(c)))  # type: ignore
        if not charset:
            _bf_print("error", error="Empty charset")
            return {"success": False, "error": "Empty charset"}
        candidates = (
            "".join(p)
            for ln in range(min_len, max_len + 1)
            for p in itertools.product(charset, repeat=ln)
        )
        total = sum(len(charset) ** ln for ln in range(min_len, max_len + 1))
    else:
        if dict_path:
            dict_path = resolve_input_path(dict_path, __file__)
        if not dict_path or not os.path.exists(dict_path):
            _bf_print("error", error="Dictionary file not found")
            return {"success": False, "error": "Dictionary file not found"}
        with open(dict_path, "r", encoding="utf-8", errors="replace") as f:
            words = [line.rstrip("\n") for line in f]
        candidates = iter(words)
        total = len(words)

    # ── Shared memory for PDF bytes ───────────────────────────────────────────
    shm = SharedMemory(create=True, size=max(1, len(pdf_bytes)))
    shm.buf[:len(pdf_bytes)] = pdf_bytes
    shm_name, shm_size = shm.name, len(pdf_bytes)

    tried = 0
    found_pwd = None
    start_time = time.time()
    last_report = 0.0

    try:
        with multiprocessing.Pool(
            processes=n_workers,
            initializer=_bf_init,
            initargs=(shm_name, shm_size),
        ) as pool:
            for result in pool.imap_unordered(
                _bf_try_chunk,
                _chunk_gen(candidates, chunk_size),
                chunksize=1,
            ):
                tried += chunk_size
                now = time.time()
                elapsed = now - start_time
                speed = int(tried / elapsed) if elapsed > 0 else 0
                eta = int((total - tried) / speed) if speed > 0 and tried < total else -1

                if now - last_report >= 0.4:
                    _bf_print("progress", tried=min(tried, total), total=total,
                              speed=speed, eta=eta, elapsed=round(elapsed, 1))
                    last_report = now

                if result is not None:
                    found_pwd = result
                    pool.terminate()
                    break
    finally:
        shm.close()
        try:
            shm.unlink()
        except Exception:
            pass

    elapsed = time.time() - start_time
    if found_pwd is not None:
        _bf_print("found", password=found_pwd, tried=min(tried, total),
                  elapsed=round(elapsed, 1))
        return {"success": True, "found": True, "password": found_pwd,
                "tried": tried, "elapsed": round(elapsed, 1)}
    else:
        _bf_print("done", found=False, tried=min(tried, total),
                  elapsed=round(elapsed, 1))
        return {"success": True, "found": False, "tried": tried,
                "elapsed": round(elapsed, 1)}


def _bf_print(event_type: str, **kwargs):
    """Print a JSON progress line to stdout (Rust reads this)."""
    print(json.dumps({"type": event_type, **kwargs}, ensure_ascii=False), flush=True)


OPERATIONS = {
    "merge": merge_pdfs,
    "split": split_pdf,
    "encrypt": encrypt_pdf,
    "decrypt": decrypt_pdf,
    "compress": compress_pdf,
    "watermark": add_watermark,
    "preview_watermark": preview_watermark,
    "page_numbers": add_page_numbers,
    "preview_pages": preview_pages,
    "preview_page_numbers": preview_page_numbers,
    "reorder_pages": reorder_pdf_pages,
    "to_docx": pdf_to_docx,
    "info": get_pdf_info,
    "bruteforce": bruteforce_pdf,
}

if __name__ == "__main__":
    # Allow bruteforce to be called with streaming output (no final JSON print)
    parser = argparse.ArgumentParser(description="PDF Handler")
    parser.add_argument("operation", choices=list(OPERATIONS.keys()))
    parser.add_argument("params", nargs="?", help="JSON-encoded parameters")
    args = parser.parse_args()

    params = json.loads(args.params) if args.params else {}
    if args.operation == "bruteforce":
        # bruteforce streams progress to stdout itself; just call it
        OPERATIONS[args.operation](**params)
    else:
        result = OPERATIONS[args.operation](**params)
        print(json.dumps(result, ensure_ascii=False))

